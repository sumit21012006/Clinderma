"""
Clinderma KB Ingestion Script — V3 (Comprehensive Multi-Source Semantic Index)

Ingests and aggregates:
  1. Master FAQs Document (DOCX)
  2. Clinical Training Module (DOCX)
  3. 51 Clinderma Website Blog Articles (JSON)
  4. Kandid AI Benchmark Responses & Product Catalog (JSON/JSONL)

Generates high-dimensional semantic embeddings using Google Gemini (gemini-embedding-001)
with automatic 429 rate limit backoff and builds a normalized FAISS vector index.
"""

import os
import sys
import json
import re
import time
import numpy as np
import docx

# Ensure Windows terminal utf-8 support
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

# Add backend to path so we can import config
sys.path.insert(0, os.path.normpath(os.path.join(os.path.dirname(__file__), "..")))
from app.core.config import settings

DATASET_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "..", "Dataset"))
DATA_DIR = settings.DATA_DIR
OUTPUT_JSON = settings.KB_INDEX_PATH
OUTPUT_FAISS = settings.FAISS_INDEX_PATH
OUTPUT_META = settings.FAISS_META_PATH


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace('\xa0', ' ').replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"').replace('\u2013', '-').replace('\u2014', '-').strip()
    return re.sub(r'\s+', ' ', text)


def find_file_in_dir(directory: str, pattern: str) -> str:
    if not os.path.exists(directory):
        return ""
    for f in os.listdir(directory):
        if pattern.lower() in f.lower():
            return os.path.join(directory, f)
    return ""


def parse_faqs_docx(file_path: str):
    if not file_path or not os.path.exists(file_path):
        print(f"[Warning] FAQ docx not found: {file_path}")
        return []

    doc = docx.Document(file_path)
    entries = []
    current_category = "General FAQs"
    current_question = None
    current_answer = []

    for para in doc.paragraphs:
        raw_text = para.text.strip()
        if not raw_text:
            continue
        text = clean_text(raw_text)

        cat_match = re.match(r'^\d+\.\s+([A-Z\s\&\-\–\']+)(?:\:)?$', text)
        if cat_match and not text.lower().startswith("is ") and not text.lower().startswith("how "):
            current_category = cat_match.group(1).title()
            continue

        if "cohort" in text.lower() or "faqs for" in text.lower():
            current_category = text.title()
            continue

        q_match = re.match(r'^(?:Q\d+[:\.]?|\d+\.)\s*(.+)$', text, re.IGNORECASE)
        if q_match:
            q_candidate = q_match.group(1).strip()

            if current_question and current_answer:
                entries.append({
                    "id": f"faq_{len(entries)+1}",
                    "source": "Master FAQs Document",
                    "category": current_category,
                    "question": current_question,
                    "answer": "\n".join(current_answer).strip(),
                    "text": f"Category: {current_category}\nQuestion: {current_question}\nAnswer: {' '.join(current_answer).strip()}"
                })
                current_question = None
                current_answer = []

            if '?' in q_candidate:
                q_parts = q_candidate.split('?', 1)
                current_question = q_parts[0].strip() + '?'
                if q_parts[1].strip():
                    current_answer.append(q_parts[1].strip())
            else:
                current_question = q_candidate
        else:
            if current_question:
                current_answer.append(text)
            else:
                if len(text) > 15:
                    entries.append({
                        "id": f"faq_{len(entries)+1}",
                        "source": "Master FAQs Document",
                        "category": current_category,
                        "question": f"Protocol / Guideline ({current_category})",
                        "answer": text,
                        "text": f"Category: {current_category}\nGuideline: {text}"
                    })

    if current_question and current_answer:
        entries.append({
            "id": f"faq_{len(entries)+1}",
            "source": "Master FAQs Document",
            "category": current_category,
            "question": current_question,
            "answer": "\n".join(current_answer).strip(),
            "text": f"Category: {current_category}\nQuestion: {current_question}\nAnswer: {' '.join(current_answer).strip()}"
        })

    return entries


def parse_module_docx(file_path: str):
    if not file_path or not os.path.exists(file_path):
        print(f"[Warning] Module docx not found: {file_path}")
        return []

    doc = docx.Document(file_path)
    entries = []
    current_section = "Clinderma Clinical Philosophy & Protocols"
    buffer = []

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        is_heading = len(text) < 70 and (text.isupper() or re.match(r'^\d+[\.\)]\s+', text) or text.endswith(':'))
        if is_heading:
            if buffer:
                content = "\n".join(buffer).strip()
                if len(content) > 25:
                    entries.append({
                        "id": f"mod_{len(entries)+1}",
                        "source": "Clinderma Training Module",
                        "category": current_section,
                        "question": f"Topic: {current_section}",
                        "answer": content,
                        "text": f"Section: {current_section}\nContent: {content}"
                    })
                buffer = []
            current_section = text
        else:
            buffer.append(text)

    if buffer:
        content = "\n".join(buffer).strip()
        if len(content) > 25:
            entries.append({
                "id": f"mod_{len(entries)+1}",
                "source": "Clinderma Training Module",
                "category": current_section,
                "question": f"Topic: {current_section}",
                "answer": content,
                "text": f"Section: {current_section}\nContent: {content}"
            })

    return entries


def parse_blogs_json(file_path: str):
    """
    Parses 51 blog articles into focused topical chunks.
    Extracts sections using markdown headings or splits long content logically.
    """
    if not file_path or not os.path.exists(file_path):
        print(f"[Warning] Blogs file not found: {file_path}")
        return []

    with open(file_path, "r", encoding="utf-8") as f:
        blogs = json.load(f)

    entries = []
    for idx, blog in enumerate(blogs, 1):
        title = clean_text(blog.get("title", f"Blog Article {idx}"))
        content = blog.get("content", "")
        if not content:
            continue

        # Split by markdown headers (e.g. **Heading** or ## Heading)
        pattern = r'(?:^|\n)(?:\*\*|#{2,3}\s*)([^\*\n#]{4,80})(?:\*\*|(?:\n|$))'
        splits = re.split(pattern, content)

        sections = []
        if len(splits) > 1:
            intro = clean_text(splits[0])
            if len(intro) > 30:
                sections.append(("Introduction & Overview", intro))
            for i in range(1, len(splits), 2):
                heading = clean_text(splits[i])
                sec_body = clean_text(splits[i+1]) if i+1 < len(splits) else ""
                if len(sec_body) > 20:
                    sections.append((heading, sec_body))
        else:
            # Paragraph splitting for unstructured articles
            paras = [clean_text(p) for p in content.split("\n\n") if clean_text(p)]
            current_chunk = []
            curr_len = 0
            chunk_num = 1
            for p in paras:
                current_chunk.append(p)
                curr_len += len(p)
                if curr_len > 800:
                    sections.append((f"Section {chunk_num}", "\n\n".join(current_chunk)))
                    current_chunk = []
                    curr_len = 0
                    chunk_num += 1
            if current_chunk:
                sections.append((f"Section {chunk_num}", "\n\n".join(current_chunk)))

        for sec_idx, (sec_heading, sec_text) in enumerate(sections, 1):
            if len(sec_text) < 15:
                continue
            entry_id = f"blog_{idx}_{sec_idx}"
            entries.append({
                "id": entry_id,
                "source": f"Clinderma Clinical Blog: {title}",
                "category": "Skincare & Clinical Insights (Blog)",
                "question": f"{title} - {sec_heading}",
                "answer": sec_text,
                "text": f"Article: {title}\nTopic: {sec_heading}\nClinical Details: {sec_text}"
            })

    return entries


def parse_kandid_results(dir_path: str):
    """
    Parses tested Kandid AI benchmark responses and product recommendations.
    """
    if not dir_path or not os.path.exists(dir_path):
        return []
    entries = []
    seen_questions = set()

    for fname in os.listdir(dir_path):
        if fname.endswith(".json") and not fname.startswith("latest"):
            fpath = os.path.join(dir_path, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    results = data.get("results", [])
                    cohort = fname.replace("FAQ_Comparision_", "").replace(".json", "").replace("_", " ")
                    for item in results:
                        q = clean_text(item.get("question", ""))
                        if not q or q.lower() in seen_questions:
                            continue
                        seen_questions.add(q.lower())

                        resp = clean_text(item.get("chatbot_response", ""))
                        if not resp:
                            continue

                        products = item.get("products_recommended", [])
                        prod_text = ""
                        if products:
                            prod_names = [p.get("title", "") for p in products if p.get("title")]
                            if prod_names:
                                prod_text = f"\nRecommended Products: {', '.join(prod_names)}"

                        entries.append({
                            "id": f"kandid_{item.get('id', len(entries)+1)}_{cohort[:4].lower().strip()}",
                            "source": f"Kandid AI Benchmark ({cohort})",
                            "category": item.get("category", "General"),
                            "question": q,
                            "answer": resp + prod_text,
                            "text": f"Category: {item.get('category', 'General')}\nQuestion: {q}\nAnswer: {resp}{prod_text}"
                        })
            except Exception as e:
                print(f"Error reading {fpath}: {e}")

    return entries


def parse_kandid_output_jsonl(file_path: str):
    """
    Parses streamed Kandid AI queries and product catalog items from all_responses.jsonl.
    """
    if not file_path or not os.path.exists(file_path):
        return []
    entries = []
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            if not line.strip():
                continue
            try:
                item = json.loads(line)
                q = clean_text(item.get("query", ""))
                resp = clean_text(item.get("full_text", ""))
                if q and resp:
                    products = item.get("products", [])
                    prod_info = []
                    for p in products:
                        p_title = p.get("title", "")
                        p_price = p.get("price", "")
                        p_url = p.get("url", "")
                        p_desc = clean_text(p.get("body_markdown", ""))
                        if p_title:
                            desc_part = f" - {p_desc}" if p_desc else ""
                            url_part = f" ({p_url})" if p_url else ""
                            prod_info.append(f"* {p_title} (INR {p_price}){desc_part}{url_part}")
                    prod_str = f"\n\nRelevant Products:\n" + "\n".join(prod_info) if prod_info else ""

                    entries.append({
                        "id": f"kandid_stream_{len(entries)+1}",
                        "source": "Kandid AI Verified Response",
                        "category": "Product Recommendations & Routine",
                        "question": q,
                        "answer": resp + prod_str,
                        "text": f"Question: {q}\nAnswer: {resp}{prod_str}"
                    })
            except Exception:
                pass
    return entries


def generate_embeddings(entries: list):
    """Generate Gemini embeddings for all KB entries with automatic 429 backoff retry."""
    from google import genai

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    embedding_dim = 3072  # gemini-embedding-001 output dimension
    vectors = []
    batch_size = 15

    print(f"Generating embeddings using {settings.EMBEDDING_MODEL} for {len(entries)} items...")

    for i in range(0, len(entries), batch_size):
        batch = entries[i:i + batch_size]
        texts = [e["text"][:2048] for e in batch]

        batch_embedded = False
        attempts = 0
        while not batch_embedded and attempts < 5:
            attempts += 1
            try:
                result = client.models.embed_content(
                    model=settings.EMBEDDING_MODEL,
                    contents=texts
                )
                for emb in result.embeddings:
                    vectors.append(emb.values)
                batch_embedded = True
            except Exception as e:
                err_msg = str(e)
                print(f"  Embedding batch {i//batch_size + 1} attempt {attempts} note: {err_msg[:80]}")
                if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg:
                    print(f"  --> Pausing 25s for free tier rate quota reset...")
                    time.sleep(25)
                else:
                    time.sleep(2)

        if not batch_embedded:
            for e_item in batch:
                item_done = False
                for _ in range(3):
                    try:
                        res_single = client.models.embed_content(
                            model=settings.EMBEDDING_MODEL,
                            contents=[e_item["text"][:2048]]
                        )
                        vectors.append(res_single.embeddings[0].values)
                        item_done = True
                        break
                    except Exception:
                        time.sleep(20)
                if not item_done:
                    vectors.append([0.0] * embedding_dim)

        if i + batch_size < len(entries):
            time.sleep(1.0)

        print(f"  Embedded {len(vectors)}/{len(entries)} entries successfully.")

    return np.array(vectors, dtype='float32')


def build_faiss_index(vectors: np.ndarray):
    """Build and save a normalized FAISS index for Cosine Similarity search."""
    import faiss

    dim = vectors.shape[1]
    faiss.normalize_L2(vectors)
    index = faiss.IndexFlatIP(dim)
    index.add(vectors)
    return index


def main():
    os.makedirs(DATA_DIR, exist_ok=True)

    faq_path = find_file_in_dir(DATASET_DIR, "MASTER FAQs")
    module_path = find_file_in_dir(DATASET_DIR, "module")
    blogs_path = os.path.join(DATA_DIR, "clinderma_all_blogs.json")
    if not os.path.exists(blogs_path):
        blogs_path = os.path.join(DATASET_DIR, "clinderma_all_blogs.json")
    kandid_dir = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "..", "WEBSCRAPPER", "comparison_results"))
    kandid_jsonl = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "..", "WEBSCRAPPER", "output", "all_responses.jsonl"))

    print(f"1. Ingesting Master FAQs from: {faq_path}")
    faq_entries = parse_faqs_docx(faq_path)
    print(f"   -> Parsed {len(faq_entries)} FAQ entries.")

    print(f"2. Ingesting Clinical Training Module from: {module_path}")
    module_entries = parse_module_docx(module_path)
    print(f"   -> Parsed {len(module_entries)} Module entries.")

    print(f"3. Ingesting 51 Website Blog Articles from: {blogs_path}")
    blog_entries = parse_blogs_json(blogs_path)
    print(f"   -> Parsed {len(blog_entries)} Blog section chunks.")

    print(f"4. Ingesting Kandid AI Benchmark Responses from: {kandid_dir}")
    kandid_entries = parse_kandid_results(kandid_dir)
    print(f"   -> Parsed {len(kandid_entries)} Kandid Q&A entries.")

    print(f"5. Ingesting Kandid AI Stream Responses from: {kandid_jsonl}")
    kandid_stream_entries = parse_kandid_output_jsonl(kandid_jsonl)
    print(f"   -> Parsed {len(kandid_stream_entries)} Stream entries.")

    all_entries = faq_entries + module_entries + blog_entries + kandid_entries + kandid_stream_entries
    print(f"\nTotal Unified Knowledge Base entries: {len(all_entries)}")

    # Save structured JSON KB
    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(all_entries, f, indent=2, ensure_ascii=False)
    print(f"Saved Unified JSON KB index: {OUTPUT_JSON} ({len(all_entries)} items)")

    # Generate embeddings and FAISS index
    vectors = generate_embeddings(all_entries)

    import faiss
    index = build_faiss_index(vectors)
    faiss.write_index(index, OUTPUT_FAISS)
    print(f"Saved FAISS index: {OUTPUT_FAISS} ({index.ntotal} vectors, dim={vectors.shape[1]})")

    # Save metadata mapping
    meta = [{"idx": i, "id": e["id"], "question": e["question"], "category": e.get("category", ""), "source": e.get("source", "")} for i, e in enumerate(all_entries)]
    with open(OUTPUT_META, "w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2, ensure_ascii=False)
    print(f"Saved FAISS metadata: {OUTPUT_META}")

    print(f"\nUnified Ingestion complete! {len(all_entries)} chunks indexed with high-dimensional Gemini embeddings.")


if __name__ == "__main__":
    main()
